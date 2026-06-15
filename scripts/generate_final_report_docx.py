from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


IMAGE_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$')
ORDERED_RE = re.compile(r'^\d+\.\s+(.+)$')


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=9, bold=bold)


def clean_inline(text: str) -> str:
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    text = text.replace('**', '')
    return text.strip()


def markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip('|').split('|')]


def is_table_separator(line: str) -> bool:
    cells = markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r':?-{3,}:?', cell or '') for cell in cells)


def resolve_image(markdown_path: Path, raw_path: str) -> Path | None:
    target = (markdown_path.parent / raw_path.strip()).resolve()
    if target.exists() and target.suffix.lower() != '.svg':
        return target
    if target.suffix.lower() == '.svg':
        png = target.with_suffix('.png')
        if png.exists():
            return png
    return target if target.exists() else None


def add_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D9E2EC')
    border.append(bottom)
    p_pr.append(border)


def add_code_block(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    paragraph.style = document.styles['No Spacing']
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run('\n'.join(lines))
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string('243B53')


def add_image(document: Document, image_path: Path, alt: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.35))
    except Exception:
        fallback = paragraph.add_run(f'[图片无法嵌入: {image_path.name}]')
        set_run_font(fallback, size=9, color='9A3412')
        return

    if alt:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(clean_inline(alt))
        set_run_font(caption_run, size=8, color='64748B')


def build_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = document.styles
    styles['Normal'].font.name = 'Microsoft YaHei'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    styles['Normal'].font.size = Pt(10.5)
    for index in range(1, 5):
        style = styles[f'Heading {index}']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.color.rgb = RGBColor.from_string('102A43')
        style.font.bold = True

    lines = markdown_path.read_text(encoding='utf-8').splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped == '---':
            add_horizontal_rule(document)
            index += 1
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = clean_inline(heading.group(2))
            paragraph = document.add_heading(text, level=level)
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue

        image = IMAGE_RE.search(stripped)
        if image:
            image_path = resolve_image(markdown_path, image.group(2))
            if image_path:
                add_image(document, image_path, image.group(1))
            else:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f'[缺失图片: {image.group(2)}]')
                set_run_font(run, size=9, color='B91C1C')
            index += 1
            continue

        if stripped.startswith('|') and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
            headers = markdown_table_row(stripped)
            rows: list[list[str]] = []
            index += 2
            while index < len(lines) and lines[index].strip().startswith('|'):
                rows.append(markdown_table_row(lines[index]))
                index += 1
            table = document.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            for col_index, header in enumerate(headers):
                set_cell_text(table.rows[0].cells[col_index], header, bold=True)
            for row in rows:
                cells = table.add_row().cells
                for col_index, value in enumerate(row[:len(headers)]):
                    set_cell_text(cells[col_index], value)
            document.add_paragraph()
            continue

        if stripped.startswith('- '):
            paragraph = document.add_paragraph(style='List Bullet')
            run = paragraph.add_run(clean_inline(stripped[2:]))
            set_run_font(run, size=10)
            index += 1
            continue

        ordered = ORDERED_RE.match(stripped)
        if ordered:
            paragraph = document.add_paragraph(style='List Number')
            run = paragraph.add_run(clean_inline(ordered.group(1)))
            set_run_font(run, size=10)
            index += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(clean_inline(stripped))
        set_run_font(run, size=10)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description='Generate the final delivery DOCX from Markdown.')
    parser.add_argument('--source', default='docs/final-delivery-report.md')
    parser.add_argument('--output', default='docs/final-delivery-report.docx')
    args = parser.parse_args()

    source = Path(args.source).resolve()
    output = Path(args.output).resolve()
    if not source.exists():
        raise SystemExit(f'Source Markdown not found: {source}')

    build_docx(source, output)
    print(f'Generated {output} from {source}')


if __name__ == '__main__':
    main()
